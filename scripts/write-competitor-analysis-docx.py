import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs")
os.makedirs(OUTPUT_DIR, exist_ok=True)
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "PHAN_TICH_CANH_TRANH_GPUVietnam.docx")

doc = Document()

# ── Page setup ──
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(2)

# ── Helper functions ──
def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0F)
    return h

def add_table_with_style(headers, rows, col_widths=None):
    """Add a styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows
    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # spacer
    return table

def add_paragraph(text, bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size or 11)
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

# ══════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════
title = doc.add_heading('Phân Tích Cạnh Tranh — GPU Cloud Trung Quốc & Đông Nam Á', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0F)

add_paragraph(
    f'Đối chiếu với kiến trúc GPUVietnam (2026-07) — Ngày lập: {datetime.date.today().strftime("%d/%m/%Y")}',
    italic=True, color=RGBColor(0x66, 0x66, 0x66), size=10
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# I. TỔNG QUAN CÁC NỀN TẢNG THEO PHÂN KHÚC
# ══════════════════════════════════════════════════════════════
add_heading('I. Tổng quan các nền tảng theo phân khúc', level=1)

# A. Trung Quốc
add_heading('A. Trung Quốc — Hệ sinh thái "AI Workshop" trưởng thành', level=2)

add_table_with_style(
    ['Nền tảng', 'Mô hình', 'Khách hàng', 'Quy mô (ước tính)'],
    [
        ['AutoDL', 'Marketplace GPU (giống Vast) + managed ComfyUI / WebUI images', 'Freelancer, sinh viên, startup AI', 'Rất lớn (#1 TQ, hàng chục nghìn GPU)'],
        ['仙宫云 (XianGongYun)', 'ComfyUI-first, cloud workstation chuyên sâu', 'Artist AI, studio', 'Trung bình, đang scale'],
        ['揽睿星舟 (Lanrui)', 'Model hub + GPU compute (Hugging Face + Replicate lai)', 'ML engineer, doanh nghiệp', 'Lớn'],
        ['矩池云 (Juchi)', 'GPU bare-metal + container', 'Researcher, training', 'Trung bình'],
        ['OpenBayes', 'MLOps platform (dataset + model + compute một chuỗi)', 'Enterprise ML team', 'Nhỏ - Trung bình'],
        ['青椒云', 'Cloud workstation đa ngành (không chỉ AI)', 'Designer, render, office', 'Lớn, đa ngành'],
        ['阿里云 PAI / 百度 AI Studio', 'Big cloud AI platform', 'Doanh nghiệp lớn', 'Rất lớn (hạ tầng Alibaba / Baidu)'],
    ],
    col_widths=[3.5, 5.5, 3.5, 4]
)

# B. Đông Nam Á & Quốc tế
add_heading('B. Đông Nam Á & Quốc tế phục vụ SEA', level=2)

add_table_with_style(
    ['Nền tảng', 'Mô hình', 'Khách hàng', 'Ghi chú'],
    [
        ['RunPod', 'Serverless + Pod (US, có máy ở SG)', 'Global, dev AI', 'Đối thủ trực tiếp nhất với GPUVietnam'],
        ['Vast.ai', 'Marketplace ngang hàng (host tự list GPU)', 'Global, giá rẻ nhất', 'GPUVietnam đang resell Vast'],
        ['Salad', 'Distributed consumer GPU', 'Edge compute, inference', 'Mô hình khác biệt'],
        ['ThinkDiffusion', 'ComfyUI / A1111 cloud', 'Artist, freelancer', 'Rất giống GPUVietnam, nhưng global'],
        ['RunDiffusion', 'ComfyUI cloud', 'Artist', 'Giống ThinkDiffusion'],
        ['Replicate', 'Model API (không workspace)', 'Dev tích hợp API', 'Không cạnh tranh trực tiếp nhưng cùng KH'],
        ['Hugging Face Spaces', 'Free / paid inference', 'Dev, demo', 'Miễn phí → khó cạnh tranh giá'],
    ],
    col_widths=[3.5, 5, 3.5, 4.5]
)

# ══════════════════════════════════════════════════════════════
# II. SO SÁNH CHI TIẾT
# ══════════════════════════════════════════════════════════════
add_heading('II. So sánh chi tiết với GPUVietnam', level=1)

# 1. Điểm giống
add_heading('1. Điểm GIỐNG — Những gì GPUVietnam đã làm đúng hướng', level=2)

add_table_with_style(
    ['Khía cạnh', 'Nền tảng TQ / SEA', 'GPUVietnam'],
    [
        ['Đóng gói ComfyUI sẵn sàng', 'AutoDL, XianGongYun, ThinkDiffusion đều làm', '✅ Docker image + workflow stock + env riêng'],
        ['Pre-installed models & workflows', 'Tất cả nền tảng ComfyUI-first đều có', '✅ download-models.sh, 5 workflow stock, 3 môi trường'],
        ['Billing theo giờ/phiên', 'AutoDL (theo phút), RunPod (theo giây)', '✅ Combo giờ + hourly ví'],
        ['Provider abstraction', 'AutoDL multi-datacenter, RunPod multi-region', '✅ GPUProvider interface, VastAdapter, blueprint Clore / Salad'],
        ['Trial miễn phí', 'AutoDL tặng credit, ThinkDiffusion trial', '✅ Trial 3h Starter'],
        ['Admin dashboard', 'Đều có', '✅ Admin duyệt đơn, KH, pricing, tặng giờ'],
        ['Auto-stop idle', 'RunPod, Vast, AutoDL', '✅ 60 phút idle'],
    ],
    col_widths=[4, 5.5, 7]
)

# 2. Điểm khác biệt
add_heading('2. Điểm KHÁC BIỆT — GPUVietnam đang đi đường riêng (lợi thế hoặc gap)', level=2)

add_table_with_style(
    ['Khía cạnh', 'Họ làm gì', 'GPUVietnam làm gì', 'Đánh giá'],
    [
        ['GPU SoT (source of truth)',
         'KHÔNG Control Plane tách biệt — GPU vừa compute vừa lưu state',
         'CP / Runtime v2.0 tách bạch: CP là SoT, GPU chỉ compute',
         '⭐ Lợi thế kiến trúc dài hạn — GPU chết không mất Project / Session'],
        ['Session Restore ≠ Job Resume',
         'Mất GPU = mất tất cả, phải làm lại từ đầu (AutoDL, XianGongYun)',
         'GPU chết → Project / Session còn, Attempt mới chạy lại',
         '⭐ Khác biệt sản phẩm lớn — không ai trong TQ / SEA làm'],
        ['Dual-run (Render An Toàn)',
         'Gần như KHÔNG AI có',
         'Job → 2 Attempt song song trên 2 GPU khác host',
         '⭐ Feature độc quyền — phù hợp thị trường GPU không ổn định'],
        ['Editor offline (không cần GPU)',
         'Luôn cần GPU để mở ComfyUI',
         'Đang xây dựng A0.5 → A1: soạn Workflow không cần GPU',
         '⭐ Lợi thế UX lớn'],
        ['Phương thức thanh toán',
         'Alipay / WeChat (TQ), Stripe / Credit Card (Global)',
         'Chuyển khoản + admin duyệt (chưa VNPay / PayOS)',
         '⚠️ Gap tạm thời — đối thủ TQ auto-pay tức thì'],
        ['Hỗ trợ khách hàng',
         'Chatbot + ticket system',
         'Zalo cá nhân',
         '⚠️ Gap — nhưng phù hợp quy mô startup'],
        ['Hạ tầng GPU',
         'AutoDL: datacenter riêng ở 10+ thành phố TQ. RunPod: Mỹ + EU + SG',
         'Resell Vast.ai (marketplace ngang hàng)',
         '⚠️ Gap lớn — Vast không ổn định, không kiểm soát phần cứng'],
        ['Multi-provider',
         'AutoDL tự có hạ tầng; một số dùng Alibaba Cloud',
         'Có abstraction nhưng mới chỉ chạy Vast',
         '⚠️ Gap — cần thêm Clore / Salad / RunPod để giảm rủi ro single-provider'],
        ['Pricing minh bạch & tự động',
         'AutoDL: bảng giá realtime theo cung cầu. RunPod: bid / spot',
         'Giá fix admin-set, manual',
         '⚠️ Có thể là gap khi scale'],
        ['GPU đa dạng',
         'TQ: A100, H100, 4090, 5090, 3090, V100, Ascend…',
         'Chỉ 3090 / 4090 / 5090',
         '⚠️ Hạn chế cho KH training / LLM'],
    ],
    col_widths=[2.8, 4.5, 4.5, 4.7]
)

# 3. Điểm vượt trội của nền tảng TQ
add_heading('3. Điểm VƯỢT TRỘI của nền tảng TQ (cần học)', level=2)

add_table_with_style(
    ['Điểm mạnh', 'Platform', 'Cách họ triển khai', 'GPUVietnam nên học'],
    [
        ['One-click ComfyUI (không cần Docker / SSH)',
         'XianGongYun, AutoDL',
         'Chọn template → 30s có ComfyUI chạy, link truy cập ngay',
         'GPUVietnam cũng đang làm, nhưng còn docker compose build --no-cache chưa xong → cần ưu tiên xong pipeline này'],
        ['Serverless / auto-scale',
         'RunPod Serverless',
         'Không cần "bật / tắt máy" — gửi job, auto provision GPU, auto destroy, tính theo giây',
         'GPUVietnam đang "manual start / stop" → B1 (Job / Attempt) là bước đi đúng hướng tới serverless'],
        ['Model Hub tích hợp',
         'OpenBayes, AutoDL, Civitai integration',
         'Tải model 1 click từ hub, cache nóng trên NAS, không cần download lại mỗi lần boot',
         'GPUVietnam download-models.sh tải lại mỗi lần boot → cần shared storage / model cache'],
        ['Bảng giá real-time theo cung cầu',
         'AutoDL',
         'Giá thay đổi theo giờ, theo datacenter, hiển thị minh bạch',
         'GPUVietnam giá fix → thiếu linh hoạt, khó cạnh tranh giá thấp'],
        ['GPU Spot / Preemptible giá rẻ',
         'AutoDL, RunPod',
         'GPU rẻ 50-70% nhưng có thể bị thu hồi',
         'GPUVietnam chưa có → cơ hội cho KH tiết kiệm'],
        ['Data pipeline tích hợp',
         'OpenBayes',
         'Dataset → Train → Model → Deploy một pipeline, không chỉ inference',
         'GPUVietnam chỉ inference (ComfyUI) → mở rộng thị trường training?'],
        ['Hệ sinh thái model + LoRA',
         'LiblibAI (đối tác AutoDL), Civitai',
         'KH tải model trực tiếp trên nền tảng, không cần upload thủ công',
         'GPUVietnam có tab Model / LoRA nhưng "Dùng ngay" = stub'],
        ['Multi-tenant GPU (nhiều KH chung GPU)',
         'RunPod Serverless',
         'Nhiều job chạy trên cùng GPU, tối ưu utilization',
         'GPUVietnam 1 user = 1 GPU → lãng phí khi scale'],
    ],
    col_widths=[3, 2.5, 5, 6]
)

# ══════════════════════════════════════════════════════════════
# III. KHUYẾN NGHỊ
# ══════════════════════════════════════════════════════════════
add_heading('III. GPUVietnam nên học gì? — Khuyến nghị có ưu tiên', level=1)

# P0
add_heading('P0 — Cần làm NGAY (blocker cho product-market fit)', level=2)

add_table_with_style(
    ['STT', 'Việc cần làm', 'Học từ ai', 'Lý do'],
    [
        ['1', 'Hoàn thiện pipeline Docker → Vast E2E', 'AutoDL, XianGongYun',
         'docker compose build --no-cache đang chạy, chưa push → KHÔNG demo được sản phẩm thật. Không có cái này, mọi thứ khác là vô nghĩa.'],
        ['2', 'Thêm ít nhất 1 provider nữa (Clore)', 'Tất cả nền tảng TQ đều multi-DC',
         'Vast single-provider = single point of failure. Clore đã có trong code (scripts diag-clore-*.mjs), cần triển khai adapter.'],
        ['3', 'Kiểm tra Gate 1 (GPU thật) PASS', '—',
         'Đã viết checklist G1–G6 trong docs; nếu chưa PASS thì không mở được A0.5 / A1'],
    ],
    col_widths=[1, 4, 3.5, 8]
)

# P1
add_heading('P1 — Cần làm trong 1-3 tháng (tạo khác biệt cạnh tranh)', level=2)

add_table_with_style(
    ['STT', 'Việc cần làm', 'Học từ ai', 'Lý do'],
    [
        ['4', 'Triển khai B1 (Job / Attempt + failover)', '— (GPUVietnam tiên phong)',
         'Architecture v2.0 đã freeze; đây là moat (hào cạnh tranh) lớn nhất. Không ai trong TQ / SEA làm CP / Runtime tách bạch như GPUVietnam.'],
        ['5', 'Triển khai A0.5 → A1 (Editor không cần GPU)', '— (GPUVietnam sáng tạo)',
         'Cho phép KH soạn Workflow hàng giờ không tốn tiền GPU. ThinkDiffusion / RunDiffusion KHÔNG có cái này.'],
        ['6', 'Model cache / shared storage', 'AutoDL, OpenBayes',
         'KH không phải tải lại SDXL 6GB mỗi lần boot → giảm boot time từ 15 phút xuống 2 phút'],
        ['7', 'VNPay / PayOS tự động', 'Nền tảng TQ dùng Alipay / WeChat auto',
         'Tăng conversion rate, giảm ops burden (admin không phải duyệt từng đơn CK)'],
    ],
    col_widths=[1, 4.5, 3.5, 7.5]
)

# P2
add_heading('P2 — Nên làm trong 3-6 tháng (scale & tối ưu)', level=2)

add_table_with_style(
    ['STT', 'Việc cần làm', 'Học từ ai'],
    [
        ['8', 'Serverless / auto-provision theo job', 'RunPod Serverless'],
        ['9', 'Spot GPU giá rẻ', 'AutoDL, RunPod'],
        ['10', 'Multi-GPU line (A100, H100 cho training)', 'AutoDL'],
        ['11', 'Model / LoRA "Dùng ngay" 1-click', 'LiblibAI, AutoDL'],
        ['12', 'Workflow marketplace (KH bán workflow cho nhau)', 'Chưa ai làm tốt → cơ hội'],
    ],
    col_widths=[1, 5.5, 10]
)

# ══════════════════════════════════════════════════════════════
# IV. SWOT
# ══════════════════════════════════════════════════════════════
add_heading('IV. Ma trận SWOT của GPUVietnam', level=1)

# Strengths / Weaknesses
add_heading('Điểm mạnh & Điểm yếu', level=2)

swot_table = doc.add_table(rows=6, cols=2)
swot_table.style = 'Light Grid Accent 1'

# Row 0: headers
swot_table.rows[0].cells[0].text = 'Strengths (Điểm mạnh)'
swot_table.rows[0].cells[1].text = 'Weaknesses (Điểm yếu)'
for cell in swot_table.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)
    set_cell_shading(cell, 'D9EAD3')

strengths = [
    '⭐ CP / Runtime v2.0 — GPU chết không mất Project',
    '⭐ Dual-run Render An Toàn — độc quyền',
    '⭐ Local: VNĐ, Zalo, chuyển khoản',
    '⭐ Kiến trúc sạch, abstraction tốt',
    '⭐ Startup linh hoạt, một người vận hành được',
]
weaknesses = [
    '⚠️ Resell Vast — không kiểm soát hạ tầng',
    '⚠️ Single provider — rủi ro cao',
    '⚠️ Chưa có auto-pay (VNPay / PayOS)',
    '⚠️ Chưa có shared model cache',
    '⚠️ Boot time còn lâu (tải model mỗi lần)',
]
for i, (s, w) in enumerate(zip(strengths, weaknesses)):
    swot_table.rows[i + 1].cells[0].text = s
    swot_table.rows[i + 1].cells[1].text = w
    for c in range(2):
        for p in swot_table.rows[i + 1].cells[c].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

doc.add_paragraph()

# Opportunities / Threats
add_heading('Cơ hội & Nguy cơ', level=2)

ot_table = doc.add_table(rows=5, cols=2)
ot_table.style = 'Light Grid Accent 1'

ot_table.rows[0].cells[0].text = 'Opportunities (Cơ hội)'
ot_table.rows[0].cells[1].text = 'Threats (Nguy cơ)'
for cell in ot_table.rows[0].cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)
    set_cell_shading(cell, 'CFE2F3')

opportunities = [
    '🟢 Thị trường VN chưa có ai làm ComfyUI cloud',
    '🟢 KH TQ bị kiểm duyệt nội dung → chảy sang VN',
    '🟢 Editor offline = moat lớn',
    '🟢 Dual-run phù hợp GPU marketplace không ổn định',
]
threats = [
    '🔴 AutoDL hoặc XianGongYun tiến vào VN',
    '🔴 RunPod mở region Việt Nam / Singapore',
    '🔴 Vast.ai thay đổi API / pricing đột ngột',
    '🔴 KH tự thuê Vast trực tiếp nếu không đủ giá trị gia tăng',
]
for i, (o, t) in enumerate(zip(opportunities, threats)):
    ot_table.rows[i + 1].cells[0].text = o
    ot_table.rows[i + 1].cells[1].text = t
    for c in range(2):
        for p in ot_table.rows[i + 1].cells[c].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# V. KẾT LUẬN
# ══════════════════════════════════════════════════════════════
add_heading('V. Kết luận', level=1)

conclusions = [
    'GPUVietnam đang đi đúng hướng kiến trúc, thậm chí còn đi trước các nền tảng TQ / SEA về mặt tách Control Plane khỏi Runtime. Đây là lợi thế cạnh tranh cực lớn nếu triển khai xong.',
    'Tuy nhiên, khoảng cách về vận hành & UX (chưa có pipeline Docker E2E, chưa multi-provider, chưa auto-pay, chưa model cache) đang là blocker chính. Nền tảng TQ mạnh về tốc độ triển khai và trải nghiệm "vào là chạy" — đó là thứ GPUVietnam cần ưu tiên bắt kịp trước khi mở rộng moat kiến trúc.',
    'Thứ tự hành động đề xuất: P0.1 (Docker E2E) → P0.2 (Clore adapter) → P0.3 (Gate 1 PASS) → P1.4 (B1 Job / Attempt) → P1.5 (A0.5 Editor offline).',
]

for text in conclusions:
    add_paragraph(text)

# ── Footer ──
doc.add_paragraph()
add_paragraph(
    f'Tài liệu nội bộ — GPUVietnam — {datetime.date.today().strftime("%d/%m/%Y")}',
    italic=True, color=RGBColor(0x99, 0x99, 0x99), size=9
)

# ── Save ──
doc.save(OUTPUT_PATH)
print(f"✅ Saved to: {OUTPUT_PATH}")